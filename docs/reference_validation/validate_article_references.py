from __future__ import annotations

import argparse
import csv
import json
import re
import time
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from difflib import SequenceMatcher


DEFAULT_DOCX = (
    Path(__file__).resolve().parents[2]
    / "documentation"
    / "article"
    / "A reproducible bibliographic landscape of AI in oncology.docx"
)
DEFAULT_OUTPUT_DIR = (
    Path(__file__).resolve().parents[2]
    / "documentation"
    / "article"
    / "reference_validation"
)
CROSSREF_API = "https://api.crossref.org/works"
USER_AGENT = "OncoTagger reference validator (mailto:bohdan.khilchevskyi@sciforce.tech)"


@dataclass
class ReferenceEntry:
    number: int
    paragraph_index: int
    text: str
    existing_doi: str | None
    existing_urls: list[str]
    crossref_doi: str | None = None
    crossref_title: str | None = None
    crossref_score: float | None = None
    title_match: float | None = None
    chosen_doi: str | None = None
    doi_url_status: str | None = None
    doi_url_final: str | None = None
    url_statuses: dict[str, str] | None = None
    validation_note: str = ""


@dataclass
class CitationToken:
    paragraph_index: int
    token: str
    expanded: list[int]
    superscript_all_chars: bool
    superscript_any_chars: bool
    context: str


def normalize_text(value: str) -> str:
    value = value.lower()
    value = re.sub(r"https?://\S+", " ", value)
    value = re.sub(r"doi:\s*10\.", " 10.", value)
    value = re.sub(r"[^a-z0-9]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def clean_doi(value: str) -> str:
    value = value.strip()
    value = re.sub(r"^doi:\s*", "", value, flags=re.I)
    value = value.rstrip(".,;)")
    return value


def doi_url(doi: str) -> str:
    return f"https://doi.org/{doi}"


def expand_citation_token(token: str) -> list[int]:
    token = token.replace(" ", "").replace("–", "-").replace("—", "-")
    values: list[int] = []
    for part in token.split(","):
        if not part:
            continue
        if "-" in part:
            left, right = part.split("-", 1)
            if left.isdigit() and right.isdigit():
                a, b = int(left), int(right)
                if a <= b:
                    values.extend(range(a, b + 1))
                else:
                    values.extend(range(a, b - 1, -1))
        elif part.isdigit():
            values.append(int(part))
    return values


def paragraph_char_flags(paragraph) -> tuple[str, list[bool]]:
    text = ""
    flags: list[bool] = []
    for run in paragraph.runs:
        sup = bool(run.font.superscript)
        for char in run.text:
            text += char
            flags.append(sup)
    return text, flags


def extract_reference_entries(doc: Document) -> tuple[int, int, list[ReferenceEntry]]:
    ref_start = None
    ref_end = len(doc.paragraphs)
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().lower() == "references":
            ref_start = i
        elif ref_start is not None and paragraph.text.strip().lower() in {"figure legends", "tables"}:
            ref_end = i
            break
    if ref_start is None:
        raise RuntimeError("Could not find a References heading.")

    entries: list[ReferenceEntry] = []
    current: ReferenceEntry | None = None
    doi_re = re.compile(r"(?:doi:\s*)?(10\.\d{4,9}/[-._;()/:A-Z0-9]+)", re.I)
    url_re = re.compile(r"https?://[^\s<>]+", re.I)

    for pi in range(ref_start + 1, ref_end):
        text = doc.paragraphs[pi].text.strip()
        if not text:
            continue
        match = re.match(r"^(\d+)\.\s+(.*)$", text)
        if match:
            if current is not None:
                entries.append(current)
            number = int(match.group(1))
            body = match.group(2).strip()
            dois = [clean_doi(x) for x in doi_re.findall(body)]
            urls = [x.rstrip(".,;") for x in url_re.findall(body)]
            current = ReferenceEntry(
                number=number,
                paragraph_index=pi,
                text=body,
                existing_doi=dois[0] if dois else None,
                existing_urls=urls,
            )
        elif current is not None:
            current.text = f"{current.text} {text}".strip()
            dois = [clean_doi(x) for x in doi_re.findall(current.text)]
            urls = [x.rstrip(".,;") for x in url_re.findall(current.text)]
            current.existing_doi = dois[0] if dois else None
            current.existing_urls = urls
    if current is not None:
        entries.append(current)
    return ref_start, ref_end, entries


def extract_citation_tokens(doc: Document, start_index: int, end_index: int) -> list[CitationToken]:
    tokens: list[CitationToken] = []
    allowed = set("0123456789,-–— ")
    for pi in range(start_index, end_index):
        text, flags = paragraph_char_flags(doc.paragraphs[pi])
        start: int | None = None
        buffer: list[str] = []

        def flush(end: int) -> None:
            nonlocal start, buffer
            if start is None or not buffer:
                start = None
                buffer = []
                return
            raw = "".join(buffer).strip(" ,")
            expanded = expand_citation_token(raw)
            if expanded:
                span_flags = flags[start:end]
                tokens.append(
                    CitationToken(
                        paragraph_index=pi,
                        token=raw.replace(" ", ""),
                        expanded=expanded,
                        superscript_all_chars=all(span_flags),
                        superscript_any_chars=any(span_flags),
                        context=text[max(0, start - 110) : end + 110],
                    )
                )
            start = None
            buffer = []

        for idx, (char, is_sup) in enumerate(zip(text, flags)):
            if is_sup and char in allowed:
                if start is None:
                    start = idx
                buffer.append(char)
            else:
                flush(idx)
        flush(len(text))

        # Fallback: catch citation-like groups that are only partially superscript.
        citation_re = re.compile(r"(?<=[A-Za-z).,;])(\d+(?:\s*[-–]\s*\d+)?(?:\s*,\s*\d+(?:\s*[-–]\s*\d+)?)*)")
        seen_spans = {(token.paragraph_index, token.token) for token in tokens}
        for match in citation_re.finditer(text):
            raw = match.group(1).replace(" ", "")
            span_flags = flags[match.start() : match.end()]
            if not any(span_flags) or (pi, raw) in seen_spans:
                continue
            expanded = expand_citation_token(raw)
            if not expanded:
                continue
            tokens.append(
                CitationToken(
                    paragraph_index=pi,
                    token=raw,
                    expanded=expanded,
                    superscript_all_chars=all(span_flags),
                    superscript_any_chars=any(span_flags),
                    context=text[max(0, match.start() - 110) : match.end() + 110],
                )
            )
    return tokens


def first_intro_index(doc: Document) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().lower() == "introduction":
            return i
    return 0


def crossref_lookup(reference: str, session: requests.Session) -> dict[str, Any] | None:
    params = {"query.bibliographic": reference, "rows": 3}
    try:
        response = session.get(CROSSREF_API, params=params, timeout=20)
        response.raise_for_status()
        data = response.json()
    except Exception as exc:
        return {"error": str(exc)}

    items = data.get("message", {}).get("items", [])
    if not items:
        return None

    ref_norm = normalize_text(reference)
    best = None
    best_match = -1.0
    for item in items:
        title = " ".join(item.get("title") or []).strip()
        title_norm = normalize_text(title)
        if not title_norm:
            match = 0.0
        elif title_norm in ref_norm:
            match = 1.0
        else:
            match = SequenceMatcher(None, title_norm, ref_norm).ratio()
        if match > best_match:
            best_match = match
            best = item
    if best is None:
        return None
    return {
        "doi": best.get("DOI"),
        "title": " ".join(best.get("title") or []).strip() or None,
        "score": best.get("score"),
        "title_match": best_match,
    }


def check_url(url: str, session: requests.Session) -> tuple[str, str | None]:
    headers = {"User-Agent": USER_AGENT}
    try:
        response = session.head(url, allow_redirects=True, timeout=20, headers=headers)
        if response.status_code in {405, 403} or response.status_code >= 500:
            response = session.get(url, allow_redirects=True, timeout=25, headers=headers, stream=True)
        status = response.status_code
        final_url = response.url
        response.close()
        if 200 <= status < 400:
            return f"ok:{status}", final_url
        if status in {401, 403}:
            return f"restricted:{status}", final_url
        return f"warning:{status}", final_url
    except Exception as exc:
        return f"error:{type(exc).__name__}: {exc}", None


def validate_references(entries: list[ReferenceEntry], use_crossref: bool) -> list[ReferenceEntry]:
    session = requests.Session()
    session.headers.update({"User-Agent": USER_AGENT})
    for entry in entries:
        if use_crossref and not entry.existing_doi and not entry.existing_urls:
            result = crossref_lookup(entry.text, session)
            if result and not result.get("error"):
                entry.crossref_doi = clean_doi(result["doi"]) if result.get("doi") else None
                entry.crossref_title = result.get("title")
                entry.crossref_score = result.get("score")
                entry.title_match = result.get("title_match")
            elif result and result.get("error"):
                entry.validation_note = f"Crossref lookup error: {result['error']}"
            time.sleep(0.12)

        confident_crossref = False
        if entry.crossref_doi and entry.title_match is not None:
            confident_crossref = entry.title_match >= 0.84 or (
                (entry.crossref_score or 0) >= 80 and entry.title_match >= 0.65
            )

        if entry.existing_doi:
            entry.chosen_doi = entry.existing_doi
        elif entry.crossref_doi and confident_crossref:
            entry.chosen_doi = entry.crossref_doi

        if entry.chosen_doi:
            status, final = check_url(doi_url(entry.chosen_doi), session)
            entry.doi_url_status = status
            entry.doi_url_final = final

        entry.url_statuses = {}
        for url in entry.existing_urls:
            status, final = check_url(url, session)
            entry.url_statuses[url] = f"{status} -> {final or ''}".strip()

        notes: list[str] = []
        if not entry.existing_doi and not entry.existing_urls and not entry.chosen_doi:
            notes.append("No DOI/URL found or confidently recovered.")
        if entry.crossref_doi and not entry.chosen_doi:
            notes.append("Crossref DOI candidate below confidence threshold; manual check recommended.")
        if entry.doi_url_status and (
            entry.doi_url_status.startswith("warning") or entry.doi_url_status.startswith("error")
        ):
            notes.append(f"DOI link status {entry.doi_url_status}.")
        if entry.url_statuses:
            bad = [
                url
                for url, status in entry.url_statuses.items()
                if not (status.startswith("ok") or status.startswith("restricted"))
            ]
            if bad:
                notes.append(f"{len(bad)} URL(s) returned non-OK/restricted status.")
        if entry.validation_note:
            notes.append(entry.validation_note)
        entry.validation_note = " ".join(notes)
    return entries


def citation_audit(tokens: list[CitationToken], ref_numbers: set[int]) -> dict[str, Any]:
    expanded_order: list[int] = []
    for token in tokens:
        expanded_order.extend(token.expanded)

    cited = set(expanded_order)
    first_seen: list[int] = []
    seen: set[int] = set()
    for number in expanded_order:
        if number not in seen:
            first_seen.append(number)
            seen.add(number)

    expected_prefix = list(range(1, len(first_seen) + 1))
    sequence_breaks = []
    for position, number in enumerate(first_seen, start=1):
        if number != position:
            sequence_breaks.append({"first_appearance_position": position, "reference_number": number})

    non_superscript = [asdict(t) for t in tokens if not t.superscript_all_chars]
    return {
        "citation_token_count": len(tokens),
        "cited_reference_count": len(cited),
        "cited_references": sorted(cited),
        "references_not_cited_before_reference_list": sorted(ref_numbers - cited),
        "citations_without_reference_list_entry": sorted(cited - ref_numbers),
        "first_seen_reference_order": first_seen,
        "expected_first_seen_prefix": expected_prefix,
        "first_seen_sequence_breaks": sequence_breaks,
        "non_fully_superscript_citation_tokens": non_superscript,
    }


def clear_paragraph(paragraph) -> None:
    p_element = paragraph._p
    for child in list(p_element):
        if child.tag != qn("w:pPr"):
            p_element.remove(child)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def rebuild_reference_paragraph(paragraph, number: int, body: str, link_targets: list[tuple[str, str]]) -> None:
    clear_paragraph(paragraph)
    paragraph.add_run(f"{number}. ")
    cursor = 0
    for visible, url in link_targets:
        index = body.find(visible, cursor)
        if index < 0:
            continue
        if index > cursor:
            paragraph.add_run(body[cursor:index])
        add_hyperlink(paragraph, visible, url)
        cursor = index + len(visible)
    if cursor < len(body):
        paragraph.add_run(body[cursor:])


def write_linked_docx(source: Path, output: Path, entries: list[ReferenceEntry]) -> None:
    doc = Document(source)
    by_number = {entry.number: entry for entry in entries}
    for entry in entries:
        paragraph = doc.paragraphs[entry.paragraph_index]
        body = entry.text
        link_targets: list[tuple[str, str]] = []

        if entry.existing_doi:
            doi_match = re.search(r"(doi:\s*)?10\.\d{4,9}/[-._;()/:A-Z0-9]+", body, flags=re.I)
            if doi_match:
                visible = doi_match.group(0).rstrip(".,;)")
                link_targets.append((visible, doi_url(entry.existing_doi)))
        elif entry.chosen_doi:
            addition = f" https://doi.org/{entry.chosen_doi}"
            if addition.strip() not in body:
                if body.endswith("."):
                    body = body[:-1] + f".{addition}"
                else:
                    body = body + addition
            link_targets.append((f"https://doi.org/{entry.chosen_doi}", doi_url(entry.chosen_doi)))

        for url in entry.existing_urls:
            link_targets.append((url, url))

        if link_targets:
            rebuild_reference_paragraph(paragraph, entry.number, body, link_targets)
        elif entry.number in by_number:
            # Keep paragraph unchanged when no confident link target exists.
            continue
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def write_outputs(
    output_dir: Path,
    docx_path: Path,
    entries: list[ReferenceEntry],
    tokens: list[CitationToken],
    audit: dict[str, Any],
    linked_docx: Path | None,
) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    rows = [asdict(entry) for entry in entries]
    csv_path = output_dir / "reference_validation_results.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()) if rows else [])
        writer.writeheader()
        for row in rows:
            row["existing_urls"] = "; ".join(row["existing_urls"] or [])
            row["url_statuses"] = json.dumps(row["url_statuses"] or {}, ensure_ascii=False)
            writer.writerow(row)

    citation_csv = output_dir / "citation_audit.csv"
    citation_rows = [asdict(token) for token in tokens]
    with citation_csv.open("w", newline="", encoding="utf-8") as handle:
        fieldnames = list(citation_rows[0].keys()) if citation_rows else [
            "paragraph_index",
            "token",
            "expanded",
            "superscript_all_chars",
            "superscript_any_chars",
            "context",
        ]
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for row in citation_rows:
            row["expanded"] = ";".join(str(x) for x in row["expanded"])
            writer.writerow(row)

    summary = {
        "timestamp_utc": datetime.now(timezone.utc).isoformat(),
        "docx_path": str(docx_path),
        "reference_count": len(entries),
        "references": rows,
        "citation_audit": audit,
        "linked_docx": str(linked_docx) if linked_docx else None,
    }
    (output_dir / "reference_validation_results.json").write_text(
        json.dumps(summary, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    bad_links = [
        entry
        for entry in entries
        if entry.validation_note
        or (
            entry.doi_url_status
            and (entry.doi_url_status.startswith("warning") or entry.doi_url_status.startswith("error"))
        )
        or any(
            not (status.startswith("ok") or status.startswith("restricted"))
            for status in (entry.url_statuses or {}).values()
        )
    ]
    no_link = [entry for entry in entries if not entry.chosen_doi and not entry.existing_urls]
    report = [
        "# Article Reference Validation Report",
        "",
        f"Source DOCX: `{docx_path}`",
        f"Generated: {summary['timestamp_utc']}",
        "",
        "## Summary",
        "",
        f"- Reference entries parsed: {len(entries)}",
        f"- Citation tokens parsed before the reference list: {audit['citation_token_count']}",
        f"- Cited references before the reference list: {audit['cited_reference_count']}",
        f"- References not cited before the reference list: {audit['references_not_cited_before_reference_list']}",
        f"- Citations without reference-list entry: {audit['citations_without_reference_list_entry']}",
        f"- First-seen sequence breaks: {audit['first_seen_sequence_breaks']}",
        f"- Non-fully superscript citation tokens: {len(audit['non_fully_superscript_citation_tokens'])}",
        f"- References without DOI/URL after lookup: {[entry.number for entry in no_link]}",
        "",
        "## Link Checks Needing Attention",
        "",
    ]
    if bad_links:
        for entry in bad_links:
            report.append(f"- Ref. {entry.number}: {entry.validation_note or 'Check link status.'}")
    else:
        report.append("- None.")
    report.extend(
        [
            "",
            "## Outputs",
            "",
            f"- CSV: `{csv_path}`",
            f"- Citation audit CSV: `{citation_csv}`",
            f"- JSON: `{output_dir / 'reference_validation_results.json'}`",
        ]
    )
    if linked_docx:
        report.append(f"- Linked DOCX copy: `{linked_docx}`")
    (output_dir / "reference_validation_report.md").write_text("\n".join(report) + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Validate manuscript references and citation numbering in a DOCX file.")
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument("--no-crossref", action="store_true", help="Skip Crossref DOI lookup for references lacking DOI/URL.")
    parser.add_argument("--write-linked-docx", action="store_true", help="Write a DOCX copy with DOI/URL hyperlinks in references.")
    args = parser.parse_args()

    docx_path = args.docx.resolve()
    output_dir = args.output_dir.resolve()
    doc = Document(docx_path)
    ref_start, ref_end, entries = extract_reference_entries(doc)
    intro_start = first_intro_index(doc)
    tokens = extract_citation_tokens(doc, intro_start, ref_start)
    ref_numbers = {entry.number for entry in entries}
    audit = citation_audit(tokens, ref_numbers)
    entries = validate_references(entries, use_crossref=not args.no_crossref)

    linked_docx = None
    if args.write_linked_docx:
        linked_docx = output_dir / f"{docx_path.stem}.reference_links_checked.docx"
        write_linked_docx(docx_path, linked_docx, entries)

    write_outputs(output_dir, docx_path, entries, tokens, audit, linked_docx)

    status = "PASS"
    warnings: list[str] = []
    if audit["first_seen_sequence_breaks"]:
        status = "WARNINGS"
        warnings.append("Reference first-appearance order is not strictly sequential.")
    if audit["non_fully_superscript_citation_tokens"]:
        status = "WARNINGS"
        warnings.append("Some citation tokens are not fully superscript.")
    if audit["citations_without_reference_list_entry"]:
        status = "FAIL"
        warnings.append("Some in-text citations have no reference-list entry.")
    if any(entry.validation_note for entry in entries):
        status = "WARNINGS" if status == "PASS" else status
        warnings.append("Some reference links/DOI candidates need attention.")

    print("Article reference validation completed.")
    print(f"Source DOCX: {docx_path}")
    print(f"References parsed: {len(entries)}")
    print(f"Citation tokens parsed: {len(tokens)}")
    print(f"References not cited before reference list: {audit['references_not_cited_before_reference_list']}")
    print(f"First-seen sequence breaks: {audit['first_seen_sequence_breaks'][:8]}")
    print(f"Non-fully superscript citation tokens: {len(audit['non_fully_superscript_citation_tokens'])}")
    print(f"Output directory: {output_dir}")
    if linked_docx:
        print(f"Linked DOCX copy: {linked_docx}")
    print(f"QC status: {status}")
    if warnings:
        print("Warnings:")
        for warning in warnings:
            print(f"- {warning}")


if __name__ == "__main__":
    main()
